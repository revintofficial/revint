from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/leadac_progress_report/LeadAC_Bugune_Kadar_Proje_Hikayesi.docx")


NAVY = RGBColor(17, 37, 63)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(80, 88, 98)
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "E8EEF5"
SOFT_GREEN = "EDF6F1"
SOFT_GOLD = "FFF7E3"
SOFT_RED = "FBEAEA"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(20, 24, 29)


def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph(p, before=0, after=8, line=1.333, align=None, keep_next=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep_next:
        fmt.keep_with_next = True


def add_para(doc, text="", size=11, bold=False, italic=False, color=BLACK, after=8, before=0, align=None):
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    set_paragraph(p, before=18 if level == 1 else 12 if level == 2 else 8, after=8 if level == 1 else 6, line=1.15, keep_next=True)
    color = BLUE if level in (1, 2) else DARK_BLUE
    size = 16 if level == 1 else 13 if level == 2 else 12
    run = p.add_run(text)
    set_run(run, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=4, line=1.208)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    if level:
        p.paragraph_format.left_indent = Inches(0.62)
    r = p.add_run(text)
    set_run(r, size=11, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph(p, before=0, after=4, line=1.208)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    r = p.add_run(text)
    set_run(r, size=11, color=BLACK)
    return p


def cell_text(cell, text, bold=False, color=BLACK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, before=0, after=0, line=1.15)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120, border_color="D8DEE8"):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths_dxa, header_fill=MID_GRAY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        cell_text(hdr[i], h, bold=True, color=NAVY, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], str(value), size=font_size)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, label, body, fill=SOFT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120, border_color="D8DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph(p, before=0, after=0, line=1.2)
    r = p.add_run(label + " ")
    set_run(r, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run(body)
    set_run(r2, size=10.5, color=BLACK)
    add_para(doc, "", after=4)


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.208
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.194)

    section.header.is_linked_to_previous = False
    h = section.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(h, before=0, after=0, line=1.0)
    r = h.add_run("LeadAC Proje Hikayesi")
    set_run(r, size=9, color=GRAY)
    r2 = h.add_run("    31 Mayıs 2026")
    set_run(r2, size=9, color=GRAY)

    f = section.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(f, before=0, after=0, line=1.0)
    rr = f.add_run("Sayfa ")
    set_run(rr, size=9, color=GRAY)
    add_field(f, "PAGE")


def cover(doc):
    add_para(doc, "LeadAC", size=12, bold=True, color=GRAY, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Bugüne Kadar Yaptığımız Her Şey", size=28, bold=True, color=NAVY, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Proje ilerleyişi, yazılan dokümanlar, beta öğrenimleri ve GitHub geçmişinden hazırlanmış kurucu raporu", size=13, italic=True, color=GRAY, after=26, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_callout(
        doc,
        "Kısa karar:",
        "LeadAC, sadece lead bulan bir araç olmaktan çıkıp, local business’a satış yapan ekiplerin hangi hesabı neden hedeflemesi gerektiğini öğrenen bir operational intelligence katmanına evrildi.",
        fill=SOFT_GREEN,
    )

    rows = [
        ("İncelenen dönem", "12 Nisan 2026 - 31 Mayıs 2026"),
        ("GitHub geçmişi", "145 commit; en yoğun gün 2 Mayıs, 27 commit"),
        ("Belge havuzu", "leadac files içinde 29 PDF, docs ve research altında aktif strateji belgeleri"),
        ("Ana müşteri kanıtı", "FineDine beta çalışmaları ve Camden / North London cafe denemeleri"),
        ("Belgenin dili", "Kod anlatımı değil; ürün, pazar, müşteri ve ilerleme anlatımı"),
    ]
    add_table(doc, ["Alan", "Özet"], rows, [2400, 6960], header_fill=SOFT_GOLD, font_size=10)

    add_para(
        doc,
        "Hazırlanma tarihi: 31 Mayıs 2026. Bu rapor, çalışma klasöründeki güncel dokümanlar, leadac files klasöründeki PDF’ler, research kayıtları ve main branch commit geçmişi üzerinden yazıldı.",
        size=9.5,
        italic=True,
        color=GRAY,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def main_body(doc):
    add_heading(doc, "1. Kısa Sonuç", 1)
    add_para(
        doc,
        "LeadAC’in bugüne kadarki hikayesi üç büyük hareketten oluşuyor. İlk hareket çalışan bir MVP çıkarmaktı: haritadan işletme bulma, web varlığını inceleme, fırsat skoru üretme ve satış temsilcisine anlaşılır bir ekran verme. İkinci hareket ürünü yalnızca liste ve mesaj üretiminden çıkarıp, lead’in arkasındaki ticari sinyali anlamaya götürdü. Üçüncü hareket ise FineDine ve restaurant-tech odağıyla birlikte geldi: ürün artık veriyi toplamaktan çok, farklı veri kaynaklarını birleştirip hangi aksiyonun işe yaradığını öğrenen bir sistem olarak düşünülüyor.",
    )
    add_para(
        doc,
        "Bu nedenle bugün LeadAC’in en önemli kazanımı tek bir özellik değil. Asıl kazanım, ürünün yönünün netleşmesi: local business’a satış yapan vertical SaaS ekipleri için hesap zekası, güvenilir analiz, satış açısı seçimi ve sonuçlardan öğrenen bir hafıza katmanı.",
    )
    add_callout(
        doc,
        "Bugünkü çekirdek tez:",
        "Google Maps, Openmart, Orbital, Clay, HubSpot, Smartlead ve benzeri araçlar veriyi veya aksiyonu taşır. LeadAC’in kazanacağı yer, bu parçaların üstünde karar, güven ve öğrenme katmanı olmaktır.",
        fill=SOFT_GREEN,
    )

    add_heading(doc, "2. Zaman Çizgisi", 1)
    rows = [
        ("12-14 Nisan", "İlk çalışan ürün", "Telefon tamircileri için başlayan lead keşfi, web inceleme, skor ve dashboard hattı kuruldu. İlk dağıtım ve veritabanı sorunları çözüldü."),
        ("18-20 Nisan", "SaaS dönüşümü", "Ürün bir demo dashboard’dan daha geniş bir B2B SaaS fikrine taşındı. Dikey sayfalar, branding, performans ve lead detay ekranı güçlendi."),
        ("22-23 Nisan", "AI Core ve SEO", "Analiz, zenginleştirme, e-posta, planlama ve SEO katmanları büyüdü. İlk güvenilirlik ve veri doğruluğu iyileştirmeleri başladı."),
        ("28 Nisan-1 Mayıs", "Operasyon ve lansman hazırlığı", "Worker dağıtımı, pipeline iyileştirmeleri, landing marketing, video planı ve dossier çalışmaları yapıldı."),
        ("2 Mayıs", "Güvenlik ve dayanıklılık günü", "Tek günde 27 commit ile tenant izolasyonu, ödeme güveni, kota, SSRF, webhook, e-posta güvenliği, rate limit ve UI kararlılığı sertleştirildi."),
        ("6-7 Mayıs", "Web overhaul ve beta", "Yeni positioning, buyer persona, pricing mantığı, FREE plan kararı, beta sprint çıktıları ve marketing analytics dashboard devreye girdi."),
        ("10-14 Mayıs", "SDR Brain ve Truth & Trust", "Lead detay yüzeyi SDR odaklı hale getirildi; kanıt, doğruluk, website doğrulama, dil güvenliği ve gözlemlenebilirlik güçlendirildi."),
        ("15-20 Mayıs", "Niş denemeleri", "Sürücü kursu, kuyumcu ve dikey kampanya denemeleri ile ürünün farklı local business segmentlerine nasıl adapte olacağı test edildi."),
        ("22 Mayıs", "Revenue intelligence pivotu", "Homepage ve positioning, local lead generation’dan revenue intelligence / local business sales diline çekildi."),
        ("29-31 Mayıs", "FineDine ve US-first strateji", "Restaurant-tech operational intelligence, native entegrasyon, compliance-first ABD planı ve AI lead analysis engine dokümanları olgunlaştı."),
    ]
    add_table(doc, ["Dönem", "Tema", "Ne oldu?"], rows, [1700, 2200, 5460], header_fill=MID_GRAY)

    add_heading(doc, "3. İlk Başlangıç: Hızlı MVP", 1)
    add_para(
        doc,
        "Başlangıçtaki fikir çok netti: belirli bir şehirde ve belirli bir nişte potansiyel işletmeleri hızlıca bulmak, web sitelerini incelemek, eksikleri görünür yapmak ve satış ekibine bir öncelik listesi vermek. İlk README ve MVP dokümanları bunu telefon tamircileri üzerinden anlatıyor. Buradaki amaç mükemmel ürünü beklemek değil, gerçek değer üreten ilk katmanı hızla ayağa kaldırmaktı.",
    )
    add_para(
        doc,
        "Gong’un kuruluş hikayesinden alınan ders de burada kendini gösteriyor: önce doğru soruyu bul, sonra çalışır bir uygulama katmanı çıkar, ardından pazarın verdiği sinyale göre derinleş. LeadAC de ilk ay içinde aynı refleksi izledi. Ürün önce liste ve skor verdi; sonra bu skorların arkasındaki kanıt, öneri ve satış hafızası ana mesele haline geldi.",
    )

    add_heading(doc, "4. Üründe Kurulan Ana Kabiliyetler", 1)
    rows = [
        ("Lead keşfi", "Kullanıcı bir bölge ve niş verdiğinde sistem Google Maps ve benzeri kaynaklardan taze işletmeler çıkarıyor."),
        ("Web ve dijital varlık incelemesi", "Her işletmenin sitesi, mobil uyumu, rezervasyon, online sipariş, güven sinyalleri, sosyal varlığı ve iletişim yolları kontrol ediliyor."),
        ("Yorum ve itibar analizi", "Yorumlardan güçlü ve zayıf taraflar, müşteri şikayetleri, satın alma sinyali olabilecek tekrar eden problemler çıkarılıyor."),
        ("Fırsat ve uyum skoru", "Her lead yalnızca var olup olmamasına göre değil, satılacak teklifle ne kadar uyumlu olduğuna göre değerlendiriliyor."),
        ("Dossier ve satış brief’i", "Satış temsilcisine tek bakışta ne söyleyeceğini, neden söyleyeceğini ve hangi kanıta dayandığını anlatan özet üretiliyor."),
        ("Mesaj ve teklif önerisi", "Lead’in gerçekten görünen eksiklerine göre ilk mesaj, teklif açısı ve paket önerisi hazırlanıyor."),
        ("Ekip ve workspace yapısı", "Tek kişilik oyuncak değil; farklı ekip üyeleri, roller, workspace’ler, planlar ve quota mantığıyla ajans / satış ekibi kullanımına hazırlandı."),
        ("Öğrenen hafıza", "Başarılı ve başarısız mesajlar, lead profilleri ve geçmiş sonuçlar sonraki önerileri iyileştirecek bir hafıza katmanına taşındı."),
    ]
    add_table(doc, ["Kabiliyet", "Bugünkü anlamı"], rows, [2500, 6860], header_fill=MID_GRAY)

    add_heading(doc, "5. Güven, Doğruluk ve Truth & Trust Katmanı", 1)
    add_para(
        doc,
        "FineDine beta testleri ve sprint raporları ürünün en kritik dersini ortaya çıkardı: satış temsilcisi ekrana güvenmezse, en güzel analiz bile işe yaramaz. Bu yüzden Mayıs ortasında yapılan Truth & Trust çalışması ürünün yönünü değiştiren temel işlerden biri oldu.",
    )
    add_bullet(doc, "Lead hakkında söylenen her iddianın arkasında görünür bir kanıt olması hedeflendi.")
    add_bullet(doc, "Yanlış siteyi gerçek site sanma riski azaltıldı; web varlığı birden fazla kaynaktan doğrulanmaya başladı.")
    add_bullet(doc, "Türkçe / İngilizce çıktı karışıklığı tek bir karar noktasına bağlandı.")
    add_bullet(doc, "Hatalar, gecikmeler ve başarısız analizler görünür hale getirildi; artık sistemin nerede bozulduğu izlenebiliyor.")
    add_bullet(doc, "Mobil lead detay ekranı ve satış temsilcisinin gerçek akışı daha ciddi ele alındı.")
    add_para(
        doc,
        "Bu çalışma, LeadAC’in “AI güzel cümle yazsın” seviyesinden “AI ancak kanıt varsa konuşsun” seviyesine geçişidir. Ürünün gelecekte restaurant-tech veya vertical SaaS ekiplerine güven verebilmesi için bu katman zorunlu hale geldi.",
    )

    add_heading(doc, "6. Beta Testlerinden Öğrenilenler", 1)
    add_para(
        doc,
        "FineDine beta çalışmaları LeadAC’in ilk gerçek müşteri laboratuvarı oldu. İlk turda 12 işletme, ikinci turda Camden ve North London tarafında 12 gerçek cafe / restaurant lead’i incelendi. Bu testler yalnızca “ürün çalışıyor mu?” sorusunu değil, “ürün nerede yanlış güven üretiyor?” sorusunu da cevapladı.",
    )
    rows = [
        ("Güçlü çalışan parçalar", "Google Maps ve derin araştırma verisi, F&B alt niş ayrımı, sosyal medya keşfi, yorum yönü, bazı leadlerde paket ve açının gerçekten değer üretmesi."),
        ("Round 1 kritik sorunları", "Instagram profilini web sitesi sanma, küçük yorum örnekleminden abartılı sonuç çıkarma, bazı yorum etiketlerinde halüsinasyon, Premium pakete fazla kayma, embedding ve erişim hataları."),
        ("Round 2 yeni sorunları", "Paket ve tier çelişkisi, aynı sinyalin ekranda iki kez görünmesi, ham kategori adlarının kullanıcıya yansıması, eski audit verisinin kalması, zincir restoranları independent gibi ele alma, expired site bağlamında yanlış opener yazma."),
        ("En önemli ürün dersi", "Tek seferlik audit yetmez. Sistem source freshness, chain / independent ayrımı, review grounding ve human review correction ile yaşayan bir intelligence loop olmalı."),
    ]
    add_table(doc, ["Alan", "Öğrenim"], rows, [2300, 7060], header_fill=SOFT_GOLD)
    add_callout(
        doc,
        "Beta gerçeği:",
        "FineDine testleri ürünün değer üretebildiğini gösterdi ama aynı zamanda güven katmanı olmadan satış ekibinin ürüne dayanamayacağını da kanıtladı. Bu, iyi bir beta sonucudur: problemi gerçek kullanıcı buldu, ürün de buna göre sertleşti.",
        fill=SOFT_GOLD,
    )

    add_heading(doc, "7. Konumlandırmanın Evrimi", 1)
    add_para(
        doc,
        "LeadAC’in positioning’i birkaç kez değişti ama bu değişim dağınık değil; pazarın daha doğru okunmasıyla adım adım netleşti. İlk konum, ajansların local business outbound için Apollo’nun bıraktığı boşluğu doldurmaktı: taze yerel veri, web audit’i, fırsat skoru ve grounded opener.",
    )
    add_para(
        doc,
        "Sonra pazar araştırması gösterdi ki local business data ve enrichment katmanı hızla kalabalıklaşıyor. Openmart, Resquared, Orbital, Clay workflow’ları, Google Maps scraper’ları ve Apify kombinasyonları ham veri üretimini ucuzlatıyor. Bu yüzden LeadAC’in savunulabilir yeri ham lead üretmek değil, bu verinin üstünde karar vermek oldu.",
    )
    add_para(
        doc,
        "22 Mayıs’taki positioning pivotu bu yüzden önemli: LeadAC artık daha çok vertical SaaS ekiplerinin local business’a satış yaparken ihtiyaç duyduğu revenue intelligence / account intelligence katmanı olarak anlatılıyor. Kısa haliyle: CRM deal’i bilir; LeadAC işletmenin arkasındaki operasyonel bağlamı bilir.",
    )

    add_heading(doc, "8. FineDine Pivotu", 1)
    add_para(
        doc,
        "FineDine, bu yeni yön için en iyi ilk kanıt alanı oldu. Çünkü restaurant-tech satışında problem yalnızca “restoran listesi bulmak” değil. Asıl problem, hangi restoranın FineDine’e gerçekten uygun olduğunu, hangi modülün satılması gerektiğini, zincir ile independent işletmenin nasıl ayrılacağını ve her sonucu sonraki aksiyona nasıl çevireceğini bilmektir.",
    )
    rows = [
        ("Yeni kategori", "Restaurant-Tech Operational Intelligence"),
        ("İlk hedef", "FineDine’in restaurant outbound motion’ını daha hızlı, daha doğru ve öğrenebilir hale getirmek"),
        ("Ana çıktı", "Account Intelligence Brief: işletme özeti, chain / location bağlamı, FineDine fit, eksik dijital modüller, pitch angle, next action, confidence ve kaynaklar"),
        ("Pilot hedefi", "30 gün içinde en az 30 FineDine-reviewed brief, kaynak güveni, yanlış chain pitch oranının düşmesi, ilk outcome döngüsünün öğrenmeye yansıması"),
        ("Neden değerli", "LeadAC için yalnızca kullanıcı değil, kategori kanıtı ve satılabilir case study zemini"),
    ]
    add_table(doc, ["Başlık", "FineDine tarafındaki anlamı"], rows, [2400, 6960], header_fill=MID_GRAY)

    add_heading(doc, "9. Entegrasyon Stratejisi", 1)
    add_para(
        doc,
        "FineDine belgeleri bir başka önemli kararı netleştiriyor: LeadAC her aracı değiştirmeye çalışmayacak. HubSpot pipeline’ı tutuyorsa HubSpot kalsın. Smartlead veya Instantly mesajı gönderiyorsa onlar kalsın. Openmart, Orbital, Clay veya Google Places veri getiriyorsa onlar kalsın. LeadAC bu kaynakların üstüne ortak hesap gerçeği, güven seviyesi, doğru aksiyon ve öğrenme katmanı koymalı.",
    )
    add_bullet(doc, "HubSpot: müşteri ve sonuç gerçeğinin tutulduğu yer.")
    add_bullet(doc, "Tek sender: ilk MVP’de Smartlead veya Instantly’den biri; ikisini aynı anda inşa etmek yerine gerçek müşteri stack’ine göre seçim.")
    add_bullet(doc, "Openmart ve Google Places: marketi doldurma ve yerel kimliği doğrulama kaynakları.")
    add_bullet(doc, "Apify: her account’a değil, yüksek potansiyelli account’lara derin araştırma.")
    add_bullet(doc, "Orbital, Clay, Resquared, Apollo: ilk günden ana ürün değil; varsa mevcut stack’ten sinyal olarak alınacak kaynaklar.")
    add_bullet(doc, "FineDine offer context: teknik connector kadar önemli, çünkü teklif mantığı olmadan analiz doğru satış açısı üretemez.")

    add_heading(doc, "10. ABD ve Compliance Dersi", 1)
    add_para(
        doc,
        "US-first strateji review’u ürünün sıradaki olgunluk seviyesini gösteriyor. Eğer LeadAC ABD restaurant-tech pazarında kullanılacaksa mesele yalnızca daha fazla entegrasyon kurmak değil. E-posta, telefon, mobil numara, opt-out, suppression, veri kaynağı kullanım hakkı ve Google Places saklama kuralları en baştan tasarlanmalı.",
    )
    rows = [
        ("Eski sıra", "Önce native entegrasyon, sonra account intelligence, sonra öğrenme."),
        ("Yeni sıra", "Önce compliance substrate, sonra CRM / sender / outcome rails, sonra onaylı veri kaynakları, sonra account intelligence ve öğrenme."),
        ("Ürün kararı", "Owner mobile gibi alanlar yüksek değerli data değil, yüksek riskli kanal olarak ele alınmalı."),
        ("GTM kararı", "Restaurant-tech yalnızca email-first değildir; field visit, call, referral, territory density ve local ecosystem de ürüne girmeli."),
    ]
    add_table(doc, ["Alan", "Çıkan karar"], rows, [2200, 7160], header_fill=SOFT_RED)

    add_heading(doc, "11. Marketing, Brand ve GTM", 1)
    add_para(
        doc,
        "Marketing tarafında yapılan iş yalnızca landing yazmak değildi. POSITIONING, MARKETING, BUYER-PERSONA, brand assets ve SEO dokümanları birlikte ürünün kime satılacağını, kime satılmayacağını, hangi kelimelerin kullanılacağını ve hangi kanalların yanlış sinyal getireceğini netleştirdi.",
    )
    add_bullet(doc, "İlk persona seti ajans sahipleri, SMMA, solo specialist, walk-in web agency, restaurant-tech BD ve local SEO agency üzerine kuruldu.")
    add_bullet(doc, "FREE planın halka açık seçenek olmaktan çıkarılması, yanlış kullanıcıyı elemek için stratejik karar olarak yazıldı.")
    add_bullet(doc, "Brand dili, hype yerine operasyonel gerçeklik, kanıt, sayı ve net mekanizma üzerine çekildi.")
    add_bullet(doc, "SEO ve programmatic sayfalar; niche, şehir, alternatif, karşılaştırma, glossary ve araç sayfalarıyla uzun vadeli talep yakalama yüzeyi olarak planlandı.")
    add_bullet(doc, "Launch video ve görsel asset çalışmaları, ürünün “lead dossiers, not contact rows” gibi net imgelerle anlatılması için hazırlandı.")

    add_heading(doc, "12. Ekip ve Çalışma Modeli", 1)
    add_para(
        doc,
        "Rol dokümanları ve Product Department Framework, LeadAC’in sadece özellik ekleyen bir ürün olmaması gerektiğini söylüyor. Çınar’ın rolü “neyi build etmeliyiz ve neye satmalıyız?” sorusunu araştırma, pazar, müşteri acısı, rakip ve messaging üzerinden filtrelemek. Product department ise pazar sinyalini kurucu kararına çevirecek evidence-based bir yapı olarak tarif edildi.",
    )
    add_para(
        doc,
        "Design Partner Program da aynı mantığı genişletiyor: amaç demo yapmak veya feature feedback toplamak değil; hangi operational signal, hangi ICP pattern’i ve hangi outbound workflow’un gerçekten revenue etkisi yarattığını öğrenmek. Bu program için ideal partner sayısı 3-5 olarak düşünülmüş; FineDine ilk ve en net partner adayı.",
    )

    add_heading(doc, "13. Bugün Elimizde Ne Var?", 1)
    rows = [
        ("Ürün çekirdeği", "Lead keşfi, audit, yorum analizi, sosyal / web sinyalleri, skor, dossier, mesaj, paket önerisi ve lead detay yüzeyi."),
        ("Satış hafızası", "Başarılı ve başarısız opener’lar, lead profilleri, review parçaları, karar ve outcome izleriyle öğrenme katmanı."),
        ("Güven altyapısı", "Kanıt gösterme, website doğrulama, dil güvenliği, workspace izolasyonu, ödeme ve kota savunmaları."),
        ("Pazar yönü", "Vertical SaaS ve restaurant-tech için operational revenue intelligence pozisyonu."),
        ("İlk kanıt", "FineDine beta, Camden / North London testleri, gerçek bug sınıfları ve iyileştirme planları."),
        ("GTM yüzeyi", "Homepage, pricing, persona, brand, SEO, launch video, design partner ve demo/pilot anlatımı."),
    ]
    add_table(doc, ["Alan", "Durum"], rows, [2300, 7060], header_fill=SOFT_GREEN)

    add_heading(doc, "14. Henüz Bitmeyen İşler", 1)
    add_para(
        doc,
        "Proje önemli bir noktaya geldi ama en kritik işler hâlâ önümüzde. Belgelerin ortak söylediği şey şu: LeadAC’in uzun vadeli değeri, tek seferlik analiz üretmekten değil, sonuçlardan öğrenen güvenilir bir account intelligence sistemine dönüşmekten gelecek.",
    )
    add_number(doc, "FineDine pilotu için compliance-first temel ve entegrasyon omurgası netleştirilmeli.")
    add_number(doc, "HubSpot ve tek sender üzerinden ilk outcome loop kurulmalı; ürün “sonucu geri alma” yeteneğini kanıtlamalı.")
    add_number(doc, "Analyze with AI deneyimi tek, görünür ve kanıtlı bir Account Intelligence Brief’e dönüşmeli.")
    add_number(doc, "Beta’da çıkan review, stale data, chain blindness ve paket çelişkisi gibi kalite sorunları ürün standardı olarak kapanmalı.")
    add_number(doc, "Design partner programı 3-5 ciddi partnerle işletilip tekrar eden pattern’ler ayrıştırılmalı.")
    add_number(doc, "Marketing dili FineDine / restaurant-tech kanıtıyla canlı örneğe bağlanmalı.")

    doc.add_page_break()
    add_heading(doc, "15. Sonraki 30 Gün İçin Öneri", 1)
    rows = [
        ("1. Hafta", "FineDine activation görüşmesi, veri / access checklist, ilk geography ve sub-niche kararı, compliance ve offer context netliği."),
        ("2. Hafta", "HubSpot + tek sender + kaynak policy temeli; 50-100 account ingest; source provenance ve suppression mantığı."),
        ("3. Hafta", "30 account intelligence brief üretimi, FineDine reviewer ile kalite işaretleme, chain / independent QA."),
        ("4. Hafta", "İlk outcome döngüsü, playbook update, pilot review, case study / anonymized proof kararı."),
    ]
    add_table(doc, ["Zaman", "Odak"], rows, [1800, 7560], header_fill=MID_GRAY)
    add_callout(
        doc,
        "Kurucu odak:",
        "Bu ayın başarı metriği daha fazla özellik değil; FineDine’in satış akışında güvenilir brief, doğru next action ve ilk öğrenen outcome döngüsünü göstermektir.",
        fill=SOFT_GREEN,
    )
    add_para(
        doc,
        "Bu planın amacı yeni bir özellik vitrini açmak değil; güven, aksiyon ve öğrenme döngüsünü gerçek satış akışı içinde kapatmaktır. Bu döngü çalıştığında sonraki büyüme kararı daha sağlam veriye dayanır.",
    )


def source_appendix(doc):
    doc.add_page_break()
    add_heading(doc, "Ek A. İncelenen Kaynak Haritası", 1)
    add_para(
        doc,
        "Aşağıdaki liste raporun arkasındaki çalışma materyalinin özet haritasıdır. Bazı belgeler birbirinin güncel veya PDF’e aktarılmış kopyasıdır; tekrar eden belgeler ayrı dosya olarak görüldü ama içerik yorumlanırken tek karar hattı içinde birleştirildi.",
    )
    rows = [
        ("Ürün ve mimari", "Product Documentation, AI-Native-Leadac, AI Lead Analysis Engine Update, AI Lead Analysis Engine plan", "Ürünün lead listesi değil, kanıtlı account intelligence brief üretmesi gerektiğini gösterdi."),
        ("Positioning", "Positioning, Positioning AI, Leadac Positioning FINAL, docs/positioning.md", "Agency outbound dilinden vertical SaaS / revenue intelligence diline geçişi gösterdi."),
        ("Market ve rakip", "Market analizi, Arastirilicak toollar, Kaan Bey Gong, Current tools max stack", "CRM, data, outreach ve revenue intelligence pazar katmanlarını; Gong, Apollo, Clay, Orbital, Openmart gibi tehdit / partnerleri haritaladı."),
        ("FineDine", "Project roadmap, LeadAC x FineDine Integration Strategy, Final US Integration, US Compliance Review", "FineDine’i ilk beachhead, native entegrasyonları ürün kanıtı ve compliance-first yaklaşımı zorunlu karar olarak konumladı."),
        ("Beta", "LEADAC AI BETA TEST ANALIZLERI serisi, beta-test reports", "Gerçek leadlerde değer ve güven sorunlarını ortaya çıkardı; chain, stale audit, review grounding ve paket tutarlılığı problemlerini görünür yaptı."),
        ("GTM ve marka", "Brand assets, New marketing document, Marketing, Buyer Persona, SEO notları, launch video planı", "Kime satılacağı, kime satılmayacağı, nasıl konuşulacağı ve hangi kanalların doğru olduğu netleşti."),
        ("Operasyon", "Sprint 1, Roller, Product department, Design Partner Program, Cinar features notion rehber", "Takım rolleri, sprint disiplini, product intelligence ve design partner çalışma modelini tarif etti."),
        ("GitHub", "145 commit, 12 Nisan-22 Mayıs arası main branch geçmişi", "Projenin hızla MVP’den güvenlik, AI core, SEO, marketing, beta ve positioning evrelerine geçtiğini doğruladı."),
    ]
    add_table(doc, ["Kategori", "Kaynaklar", "Rapora etkisi"], rows, [1800, 3350, 4210], header_fill=MID_GRAY, font_size=8.8)

    add_heading(doc, "Ek B. leadac files PDF Envanteri", 1)
    pdf_rows = [
        ("AI-Native-Leadac (2)/(3)", "AI-native vizyon, dynamic planning ve outcome hafızası"),
        ("Arastirilicak toollar", "Rakip, adjacent oyuncu, vertical CRM ve partner araştırma listesi"),
        ("Brand asseets", "Brand ladder, Gong savunması, kanal ve keyword stratejisi"),
        ("Cinar - features notion rehber", "Hedef, outcome ve ürün araştırma düzeni"),
        ("Current tools max stack (FineDine)", "FineDine’in maksimum modern tool stack ile nereye kadar gidebileceği"),
        ("Ilk Versiyon MVP ve Pazara Çıkış", "Gong’dan MVP ve hızlı pazar dersi"),
        ("Kaan Bey Gong", "Revenue intelligence pazarını ve Gong’un neyi sahiplendiğini anlatan referans"),
        ("KAAN BEYIN DIKKATINE SEO", "SEO ve category arama odağı"),
        ("LEADAC AI BETA TEST ANALIZLERI 12-17", "Tekil cafe / restaurant lead analizleri"),
        ("Leadac Design Partner Program", "3-5 partnerle operational memory validation modeli"),
        ("Leadac Product department", "Market sinyalini ürün kararına çeviren product intelligence fonksiyonu"),
        ("Leadac Positioning FINAL", "Operational revenue intelligence kararının güncel özeti"),
        ("LeadAC x FineDine Integration Strategy", "FineDine için native-first integration planı"),
        ("LeadAC AI Lead Analysis Engine Update", "Analyze with AI ve tek Account Intelligence Brief kararı"),
        ("LeadAC-Gelir-Zekası-Katmanı", "Apollo, Gong ve Clay güçlerini yerel pazara uyarlama düşüncesi"),
        ("Market analizi", "Sales tech pazar katmanları ve revenue intelligence trendi"),
        ("New marketing document B2B", "Product marketing context ve homepage dili"),
        ("Operational Revenue Intelligence for SMB Markets Özet", "Pivotun kısa araştırma özeti"),
        ("Positioning / Positioning AI", "Post-pivot positioning ve geniş workbook"),
        ("Product Documentation", "Koddan çıkarılmış gerçek ürün kabiliyetleri"),
        ("Project roadmap", "FineDine pivot master project paper"),
        ("Roller", "Çınar ve ekip sorumluluklarının sade tanımı"),
        ("Sprint 1", "Truth & Trust sprint çıktıları"),
    ]
    add_table(doc, ["Dosya grubu", "Ne kattı?"], pdf_rows, [3600, 5760], header_fill=SOFT_GOLD, font_size=8.7)
    add_para(
        doc,
        "Not: leadac files içindeki slack gunsonular.txt dosyası okunabilir içerik taşımadığı için ana anlatıya katkı vermedi.",
        size=9,
        italic=True,
        color=GRAY,
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    cover(doc)
    main_body(doc)
    source_appendix(doc)
    doc.save(OUT)
    print(str(OUT.resolve()))


if __name__ == "__main__":
    build()
